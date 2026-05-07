"""Renders the Sovrintendenza DOCX from a SurveillancePayload + photo bytes.

Stateless: no DB, no filesystem-of-record. Receives the snapshot, writes the
rendered DOCX into a per-call working directory, returns the path.

The structure follows the canonical sezioni di una sorveglianza archeologica:
  1. Premessa
  2. Inquadramento territoriale
  3. Inquadramento storico-archeologico
  4. Inquadramento geologico e geomorfologico
  5. Metodologia
  6. Giornale di scavo / assistenza
  7. Risultati della sorveglianza
  8. Documentazione fotografica
  9. Conclusioni
"""
from __future__ import annotations

from datetime import date as _date
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

from app.config import settings
from app.payloads import DayLogPayload, PresencePayload, SurveillancePayload


_ITALIAN_MONTHS = [
    "gennaio", "febbraio", "marzo", "aprile", "maggio", "giugno",
    "luglio", "agosto", "settembre", "ottobre", "novembre", "dicembre",
]
_ITALIAN_WEEKDAYS = [
    "lunedì", "martedì", "mercoledì", "giovedì", "venerdì", "sabato", "domenica",
]
_ROLE_LABELS = {
    "direttore_tecnico": "Direttore tecnico",
    "archeologo": "Archeologo",
    "collaboratore": "Collaboratore",
    "operatore": "Operatore",
    "rilevatore": "Rilevatore",
    "altro": "Altro",
}


def _parse_iso_date(s: str) -> _date | None:
    try:
        return _date.fromisoformat(s)
    except (TypeError, ValueError):
        return None


def _format_italian_date(iso: str) -> str:
    d = _parse_iso_date(iso)
    if d is None:
        return iso
    return f"{d.day} {_ITALIAN_MONTHS[d.month - 1]} {d.year}"


def _italian_weekday(iso: str) -> str:
    d = _parse_iso_date(iso)
    if d is None:
        return ""
    return _ITALIAN_WEEKDAYS[d.weekday()]


def _role_label(role: str | None) -> str:
    if not role:
        return "—"
    return _ROLE_LABELS.get(role, role.replace("_", " ").capitalize())


def _heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def _para(doc, text: str | None, italic: bool = False) -> None:
    if not text:
        text = "—"
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)


def _kv_table(doc, rows: list[tuple[str, str | None]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    for i, (k, v) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        c0.text = k
        c1.text = v or "—"
        for run in c0.paragraphs[0].runs:
            run.bold = True
        c0.width = Cm(5.5)
        c1.width = Cm(11.0)


def render_docx(
    s: SurveillancePayload,
    photos: dict[str, bytes],
    map_bytes: bytes | None = None,
) -> Path:
    """Render a Sovrintendenza DOCX. `photos` is keyed by the photo id;
    `map_bytes`, when present, is rendered as the tavola d'insieme di
    posizionamento topografico in section 2."""
    doc = Document()

    title = doc.add_heading("RELAZIONE DI SORVEGLIANZA ARCHEOLOGICA", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(s.title)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(14)

    doc.add_paragraph()

    _kv_table(
        doc,
        [
            ("Protocollo", s.protocollo),
            ("Committente", s.committente),
            ("Direttore tecnico", s.direttore_tecnico),
            ("Soprintendenza competente", s.sabap),
            ("Comune", s.comune),
            ("Provincia", s.provincia),
            ("Foglio catastale", s.foglio_catastale),
            ("Particelle", s.particelle),
            ("Periodo di indagine", f"{s.start_date or '—'} – {s.end_date or '—'}"),
        ],
    )

    doc.add_page_break()

    _heading(doc, "1. Premessa")
    _para(
        doc,
        s.premessa
        or (
            "La presente relazione documenta le attività di sorveglianza archeologica "
            "condotte ai sensi del D.Lgs. 42/2004 e del D.Lgs. 36/2023, su prescrizione "
            "della Soprintendenza Archeologia, Belle Arti e Paesaggio competente per "
            "territorio."
        ),
    )

    _heading(doc, "2. Inquadramento territoriale")
    _para(
        doc,
        f"L'area oggetto di intervento è ubicata nel comune di "
        f"{s.comune or '—'} (provincia di {s.provincia or '—'}), foglio catastale "
        f"{s.foglio_catastale or '—'}, particelle {s.particelle or '—'}.",
    )
    if map_bytes:
        _heading(doc, "2.1 Tavola d'insieme di posizionamento topografico", level=2)
        try:
            doc.add_picture(BytesIO(map_bytes), width=Cm(15))
            cap = doc.add_paragraph(
                f"Tavola d'insieme · {s.comune or '—'}{', ' + s.provincia if s.provincia else ''}",
            )
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.runs[0].italic = True
            cap.runs[0].font.size = Pt(10)
        except Exception:
            _para(doc, "[Tavola d'insieme non leggibile.]", italic=True)
    else:
        _para(
            doc,
            "[Tavola d'insieme: usare il pulsante 'Esporta come immagine' nella "
            "pagina mappa per allegare automaticamente la cartografia di "
            "posizionamento topografico.]",
            italic=True,
        )

    _heading(doc, "3. Inquadramento storico-archeologico")
    _para(
        doc,
        "[Sintesi del contesto archeologico noto, vincoli archeologici e "
        "monumentali, evidenze pregresse. Estratto dalla consultazione di "
        "Vincoli in Rete e dalla bibliografia di riferimento.]",
        italic=True,
    )

    _heading(doc, "4. Inquadramento geologico e geomorfologico")
    _para(
        doc,
        "[Estratto dalla Carta Geologica d'Italia (CARG) e dai dati ISPRA. "
        "Descrizione del substrato e delle dinamiche geomorfologiche rilevanti.]",
        italic=True,
    )

    _heading(doc, "5. Metodologia")
    _para(doc, s.metodologia)

    _render_giornale(doc, s)

    _heading(doc, "7. Risultati della sorveglianza")
    _para(doc, s.risultati)

    if s.findings:
        _heading(doc, "7.1 Evidenze archeologiche", level=2)
        for i, f in enumerate(s.findings, 1):
            _heading(doc, f"Evidenza {i}: {f.name}", level=3)
            _para(doc, f.description)
            if f.interpretation:
                _para(doc, f"Interpretazione: {f.interpretation}")
            if f.start_date or f.end_date:
                _para(doc, f"Datazione: {f.start_date or '—'} / {f.end_date or '—'}")
            if f.recorded_on:
                _para(doc, f"Data di rilievo: {_format_italian_date(f.recorded_on)}", italic=True)
            if f.units:
                _heading(doc, "Unità Stratigrafiche", level=4)
                for u in f.units:
                    _para(
                        doc,
                        f"US {u.number} ({u.type or '—'}): "
                        f"{u.definition or '—'} — {u.description or ''}",
                    )

    _heading(doc, "8. Documentazione fotografica")
    if s.photos:
        for p in s.photos:
            blob = photos.get(p.id)
            if blob is None:
                _para(doc, f"[immagine non trasmessa: {p.filename}]", italic=True)
                continue
            try:
                doc.add_picture(BytesIO(blob), width=Cm(14))
                cap = doc.add_paragraph(p.caption or p.filename)
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.runs[0].italic = True
            except Exception:
                _para(doc, f"[immagine non leggibile: {p.filename}]", italic=True)
    else:
        _para(doc, "[Nessuna documentazione fotografica acquisita.]", italic=True)

    _heading(doc, "9. Conclusioni")
    _para(doc, s.conclusioni)

    out_dir = settings.work_dir / f"render-{uuid4().hex}"
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path = out_dir / "sorveglianza.docx"
    doc.save(out_path)
    return out_path


def _render_giornale(doc, s: SurveillancePayload) -> None:
    """Render section 6 — Giornale di scavo / assistenza.

    Per the Sovrintendenza requirements: chronological list of site days,
    each with hourly presences, personnel, operations, and localization.
    """
    _heading(doc, "6. Giornale di scavo")
    if not s.days:
        _para(
            doc,
            "[Nessuna giornata di assistenza ancora registrata. Le giornate vengono "
            "compilate dal campo tramite l'interfaccia archaeo-pro e includono presenze, "
            "operazioni svolte e localizzazione.]",
            italic=True,
        )
        return

    sorted_days = sorted(s.days, key=lambda d: d.date)
    for day in sorted_days:
        _render_day(doc, day)


def _render_day(doc, day: DayLogPayload) -> None:
    date_label = _format_italian_date(day.date)
    weekday = _italian_weekday(day.date)
    title = f"{date_label} — {weekday}" if weekday else date_label
    _heading(doc, title, level=3)

    if day.localizzazione:
        _para(doc, f"Localizzazione del lavoro: {day.localizzazione}", italic=True)
    if day.weather:
        _para(doc, f"Condizioni meteo: {day.weather}", italic=True)

    if day.presenze:
        _render_presenze_table(doc, day.presenze)
    else:
        _para(doc, "[Nessuna presenza registrata per la giornata.]", italic=True)

    if day.operazioni:
        _heading(doc, "Operazioni svolte", level=4)
        _para(doc, day.operazioni)

    if day.notes:
        _heading(doc, "Note", level=4)
        _para(doc, day.notes, italic=True)


def _render_presenze_table(doc, presenze: list[PresencePayload]) -> None:
    table = doc.add_table(rows=len(presenze) + 1, cols=5)
    table.style = "Table Grid"
    table.autofit = False

    headers = ("Nome", "Ruolo", "Inizio", "Fine", "Ore")
    widths = (Cm(5.5), Cm(4.0), Cm(2.0), Cm(2.0), Cm(1.5))
    hdr = table.rows[0].cells
    for i, label in enumerate(headers):
        hdr[i].text = label
        hdr[i].width = widths[i]
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)

    total_hours = 0.0
    for i, p in enumerate(presenze, 1):
        row = table.rows[i].cells
        row[0].text = p.name or "—"
        row[1].text = _role_label(p.role)
        row[2].text = p.hours_start or "—"
        row[3].text = p.hours_end or "—"
        if p.hours_total is not None:
            row[4].text = f"{p.hours_total:g}"
            total_hours += p.hours_total
        else:
            row[4].text = "—"
        for cell, w in zip(row, widths):
            cell.width = w
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)

    if total_hours:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Totale ore della giornata: {total_hours:g}")
        run.italic = True
        run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(8)
